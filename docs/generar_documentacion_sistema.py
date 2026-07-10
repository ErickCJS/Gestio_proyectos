from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Documentacion_tecnica_MAGERISK.docx")


CITAS = {
    "C1": "requirements.txt:4-20. Dependencias principales: FastAPI, Jinja2, PyMySQL, bcrypt, Starlette, Uvicorn y python-multipart.",
    "C2": "app/main.py:7-11, 716-730. Inicializacion de FastAPI, plantillas Jinja2, archivos estaticos, sesiones y ruta raiz.",
    "C3": "app/conexion.py:1-13. Conexion a MySQL mediante PyMySQL, base gestion_riesgo y cursor tipo diccionario.",
    "C4": "app/routes/usuario.py:36-217. Registro de usuarios, validaciones de formulario, verificacion de correo existente y hash con bcrypt.",
    "C5": "app/routes/procesos.py:29-70, 132-143, 356-391, 426-515. Tablas de roles, integrantes, procesos, grupos y asignacion de integrantes.",
    "C6": "app/routes/riesgos.py:20-49, 97-121, 137-211, 228-448. Matriz de impacto/probabilidad, CRUD de riesgos y asociacion con procesos.",
    "C7": "app/routes/controles.py:19-25, 26-93, 119-241. Validacion porcentual, carga de riesgos/controles y CRUD de controles.",
    "C8": "app/main.py:93-199. Calculo de riesgo inherente y residual con reducciones por controles.",
    "C9": "app/main.py:201-421. Consolidacion de metricas del dashboard, procesos, grupos, responsables, riesgos y controles.",
    "C10": "app/main.py:424-715, 747-815. Generacion manual de archivos XLSX y endpoints de exportacion.",
    "C11": "app/templates/maestras/dashboard.html:147-260. Estructura visual del dashboard ejecutivo y entrega de datos al JavaScript.",
    "C12": "static/js/dashboard_interactivo.js:60-183, 233-488. Construccion de mapas de calor, filtros, detalle y exportacion a imagen.",
    "C13": "static/js/librerias.js:3-41. Inicializacion de DataTables con idioma espanol y configuracion de columnas.",
    "C14": "app/templates/riesgo.html:52-276, app/templates/procesos.html:28-156, app/templates/controles.html:24-187. Plantillas de listado, detalle y acciones por modulo.",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def add_paragraph(doc, text="", style=None, citation=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    p.add_run(text)
    if citation:
        r = p.add_run(" [{}]".format(citation))
        r.font.color.rgb = RGBColor(31, 78, 121)
        r.bold = True
    return p


def add_bullet(doc, text, citation=None):
    p = add_paragraph(doc, text, style="List Bullet", citation=citation)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.1)
    return p


def add_number(doc, text, citation=None):
    p = add_paragraph(doc, text, style="List Number", citation=citation)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def add_callout(doc, title, body, citation=None):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run("\n" + body)
    if citation:
        r = p.add_run(" [{}]".format(citation))
        r.font.color.rgb = RGBColor(31, 78, 121)
        r.bold = True
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True, color="0B2545")
        if widths:
            cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level > 1 else 14)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181 if level <= 2 else 78)
    return p


def build_document():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Documentacion tecnica y metodologica del sistema MAGERISK")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Gestion de riesgos, procesos, grupos, controles, dashboard y exportaciones").italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Generado el {} | Proyecto: Gestio_proyectos".format(date.today().isoformat()))

    add_callout(
        doc,
        "Proposito del documento",
        "Este informe documenta la resolucion del sistema completo: que se construyo, por que se decidio asi, como se organiza tecnicamente, que metodologia de gestion de riesgos aplica y donde se evidencia cada afirmacion en el codigo fuente.",
    )

    add_heading(doc, "1. Resumen ejecutivo")
    add_paragraph(
        doc,
        "MAGERISK es una aplicacion web orientada a registrar grupos, procesos, riesgos y controles; calcular niveles de exposicion; visualizar mapas de calor; y exportar informacion ejecutiva para analisis posterior. La solucion usa FastAPI como backend, Jinja2 como motor de plantillas, MySQL como persistencia y JavaScript para la interaccion del dashboard.",
        "Normal",
        "C1",
    )
    add_paragraph(
        doc,
        "La estructura responde a una metodologia por capas: primero se define la informacion organizacional (grupos y procesos), luego se registran riesgos por impacto y probabilidad, despues se asocian controles, y finalmente se consolidan metricas ejecutivas para priorizar decisiones.",
        "Normal",
        "C5",
    )

    add_heading(doc, "2. Metodologia aplicada")
    add_paragraph(
        doc,
        "La metodologia del sistema se basa en una lectura clasica de gestion de riesgos: identificar el riesgo, estimar su probabilidad e impacto, ubicarlo en una matriz, aplicar controles y comparar la exposicion inherente contra la residual.",
        "Normal",
        "C6",
    )
    add_number(doc, "Levantamiento de entidades: grupos, integrantes, roles, procesos, riesgos y controles.", "C5")
    add_number(doc, "Evaluacion inherente: cada riesgo se clasifica por impacto y probabilidad antes de considerar controles.", "C6")
    add_number(doc, "Tratamiento del riesgo: los controles reducen probabilidad, impacto o ambos segun solidez y mitigacion declarada.", "C7")
    add_number(doc, "Consolidacion ejecutiva: el dashboard integra conteos, niveles, responsables, procesos relacionados y mapas de calor.", "C9")
    add_number(doc, "Salida documentaria/operativa: se generan exportaciones XLSX para dashboard completo, proceso filtrado o riesgo seleccionado.", "C10")

    add_callout(
        doc,
        "Criterio de diseno metodologico",
        "La razon de separar riesgo inherente y residual es permitir que el usuario vea la exposicion inicial y la exposicion despues de los controles. Esta comparacion evita que el dashboard sea solo un inventario y lo convierte en una herramienta de priorizacion.",
        "C8",
    )

    add_heading(doc, "3. Arquitectura general")
    add_table(
        doc,
        ["Capa", "Responsabilidad", "Evidencia"],
        [
            ["Backend", "Define la app FastAPI, rutas principales, sesiones, plantillas y exportaciones.", "[C2]"],
            ["Persistencia", "Centraliza la conexion a MySQL y devuelve resultados como diccionarios.", "[C3]"],
            ["Rutas de dominio", "Agrupa comportamiento por usuarios, procesos/grupos, riesgos y controles.", "[C4], [C5], [C6], [C7]"],
            ["Vista", "Usa plantillas Jinja2 para tablas, formularios, dashboard y detalle.", "[C11], [C14]"],
            ["Interaccion", "Renderiza mapa de calor, filtros, detalles y exportacion de imagen en cliente.", "[C12]"],
            ["Tablas dinamicas", "Aplica DataTables con idioma espanol a los listados.", "[C13]"],
        ],
        widths=[1.35, 3.8, 1.15],
    )
    add_paragraph(
        doc,
        "El patron dominante es una aplicacion server-rendered: el servidor arma el contexto, Jinja2 lo entrega como HTML y, para el dashboard, se serializa informacion consolidada hacia JavaScript mediante una variable global.",
        "Normal",
        "C11",
    )

    add_heading(doc, "4. Modulos funcionales")
    add_heading(doc, "4.1 Usuarios", level=2)
    add_paragraph(
        doc,
        "El modulo de usuarios implementa registro con validaciones basicas: campos obligatorios, longitud minima del nombre, formato de correo, complejidad minima de contrasena, verificacion de correo duplicado y almacenamiento de la contrasena con hash bcrypt.",
        "Normal",
        "C4",
    )
    add_bullet(doc, "Por que bcrypt: porque no guarda la contrasena en texto plano y genera un hash preparado para autenticacion segura.", "C4")
    add_bullet(doc, "Por que validar antes de insertar: porque reduce datos invalidos y evita registros incompletos en la base.", "C4")

    add_heading(doc, "4.2 Grupos, roles, integrantes y procesos", level=2)
    add_paragraph(
        doc,
        "Los grupos representan unidades responsables. Los procesos se asocian a un grupo, y los integrantes se vinculan con un rol activo o inactivo. Esta decision permite saber no solo donde ocurre un riesgo, sino tambien quien tiene responsabilidad organizacional sobre el proceso.",
        "Normal",
        "C5",
    )
    add_table(
        doc,
        ["Entidad", "Uso en el sistema", "Justificacion"],
        [
            ["grupo", "Agrupa procesos e integrantes.", "Permite atribuir responsabilidades."],
            ["rol", "Define el papel del integrante.", "Ordena la participacion por responsabilidades."],
            ["usuario_grupo", "Relaciona usuario, grupo y rol.", "Evita duplicidad con restriccion unica usuario-grupo."],
            ["proceso", "Representa flujo o actividad organizacional.", "Conecta la operacion con el riesgo."],
        ],
        widths=[1.35, 2.35, 2.6],
    )

    add_heading(doc, "4.3 Riesgos", level=2)
    add_paragraph(
        doc,
        "El modulo de riesgos aplica una matriz de 5x5. Impacto y probabilidad se convierten en valores de 1 a 5, se multiplican y se clasifican como MUY BAJO, BAJO, MEDIO, ALTO o EXTREMO.",
        "Normal",
        "C6",
    )
    add_table(
        doc,
        ["Puntaje", "Nivel resultante", "Razon metodologica"],
        [
            ["1", "MUY BAJO", "Exposicion minima: baja urgencia de accion."],
            ["2 a 4", "BAJO", "Riesgo controlable con seguimiento regular."],
            ["5 a 9", "MEDIO", "Requiere revision y priorizacion."],
            ["10 a 16", "ALTO", "Exige tratamiento y controles visibles."],
            ["17 a 25", "EXTREMO", "Prioridad maxima por exposicion critica."],
        ],
        widths=[1.1, 1.55, 3.65],
    )
    add_paragraph(
        doc,
        "La asociacion riesgo-proceso se maneja con una tabla intermedia, lo cual permite que un riesgo pertenezca a varios procesos y que un proceso tenga varios riesgos.",
        "Normal",
        "C6",
    )

    add_heading(doc, "4.4 Controles", level=2)
    add_paragraph(
        doc,
        "Los controles se asocian a un riesgo, tienen tipo, solidez y porcentajes de mitigacion sobre probabilidad e impacto. La funcion de validacion limita los porcentajes entre 0 y 100 para evitar valores fuera de rango.",
        "Normal",
        "C7",
    )
    add_table(
        doc,
        ["Campo/criterio", "Funcion", "Por que existe"],
        [
            ["tipo", "Preventivo, Detectivo o Correctivo.", "Diferencia la naturaleza del control."],
            ["solidez_control", "Muy baja a Muy alta.", "Pondera la capacidad real del control."],
            ["mitigacion_probabilidad", "Porcentaje de reduccion de frecuencia.", "Modela controles que evitan ocurrencia."],
            ["mitigacion_impacto", "Porcentaje de reduccion de consecuencias.", "Modela controles que disminuyen dano."],
        ],
        widths=[1.75, 2.25, 2.3],
    )

    add_heading(doc, "5. Calculo del riesgo")
    add_heading(doc, "5.1 Riesgo inherente", level=2)
    add_paragraph(
        doc,
        "El riesgo inherente se calcula transformando probabilidad e impacto a porcentajes de 20, 40, 60, 80 o 100. Luego se multiplica probabilidad por impacto y se divide entre 100. Esto produce un valor comparable antes de aplicar controles.",
        "Normal",
        "C8",
    )
    add_paragraph(doc, "Formula: Riesgo inherente = (probabilidad inicial x impacto inicial) / 100.")

    add_heading(doc, "5.2 Riesgo residual", level=2)
    add_paragraph(
        doc,
        "El riesgo residual incorpora la solidez del control y los porcentajes de mitigacion. Para cada control se calcula una capacidad real y una reduccion real. Luego se promedian las reducciones aplicables y se restan de los valores iniciales de probabilidad e impacto.",
        "Normal",
        "C8",
    )
    add_paragraph(doc, "Formula de reduccion: reduccion = (maximo_baja x solidez x mitigacion) / 10000.")
    add_paragraph(doc, "Formula residual: riesgo residual = (probabilidad residual x impacto residual) / 100.")
    add_callout(
        doc,
        "Por que usar promedio de reducciones",
        "El sistema toma las reducciones de controles que realmente mitigan probabilidad o impacto y calcula un promedio. Asi evita sumar controles indefinidamente hasta producir reducciones artificiales mayores a la exposicion original.",
        "C8",
    )

    add_heading(doc, "6. Dashboard y analitica visual")
    add_paragraph(
        doc,
        "El dashboard consolida conteos de grupos, procesos, riesgos y controles; distribucion por nivel; detalle de procesos con grupos e integrantes; y detalle de riesgos con procesos, grupos, controles y calculos inherentes/residuales.",
        "Normal",
        "C9",
    )
    add_paragraph(
        doc,
        "La plantilla del dashboard expone estos datos al cliente como window.dashboardRiesgosData, y el JavaScript los usa para construir filtros por proceso, mapas de calor, lista priorizada de riesgos y panel de detalle.",
        "Normal",
        "C11",
    )
    add_bullet(doc, "Mapa inherente: usa impacto y probabilidad originales del riesgo.", "C12")
    add_bullet(doc, "Mapa residual: usa impacto y probabilidad categorizados despues de controles.", "C12")
    add_bullet(doc, "Exportacion a imagen: genera un canvas comparativo de mapas inherente y residual.", "C12")

    add_heading(doc, "7. Exportaciones")
    add_paragraph(
        doc,
        "El sistema genera archivos XLSX directamente mediante XML empaquetado en ZIP. Existen exportaciones para dashboard completo, dashboard filtrado por proceso y detalle de un riesgo seleccionado.",
        "Normal",
        "C10",
    )
    add_table(
        doc,
        ["Exportacion", "Endpoint", "Contenido"],
        [
            ["Dashboard", "/dashboard/exportar_excel", "Resumen, niveles y detalle consolidado de riesgos."],
            ["Dashboard por proceso", "/dashboard/exportar_excel?id_proceso=...", "Riesgos asociados a un proceso."],
            ["Detalle de riesgo", "/dashboard/riesgo/{id_riesgo}/exportar_excel", "Resumen, mapa, controles y reducciones del riesgo."],
        ],
        widths=[1.55, 2.6, 2.15],
    )

    add_heading(doc, "8. Seguridad y control de datos")
    add_paragraph(
        doc,
        "La aplicacion utiliza sesiones de Starlette para mensajes y estado de usuario, contrasenas con bcrypt durante el registro, consultas parametrizadas en los INSERT/UPDATE/DELETE principales y validaciones de formulario antes de escribir en base de datos.",
        "Normal",
        "C2",
    )
    add_bullet(doc, "Fortaleza: uso de bcrypt para hash de contrasenas.", "C4")
    add_bullet(doc, "Fortaleza: uso de parametros SQL en operaciones de escritura y consulta puntual.", "C4")
    add_bullet(doc, "Observacion: la clave secreta de sesiones esta escrita directamente en codigo y deberia moverse a variables de entorno.", "C2")
    add_bullet(doc, "Observacion: las credenciales de MySQL estan fijas en conexion.py y deberian externalizarse en .env.", "C3")
    add_bullet(doc, "Observacion: el login actual redirige al dashboard sin verificacion real de credenciales; queda como punto pendiente funcional.", "C4")

    add_heading(doc, "9. Modelo de datos inferido")
    add_table(
        doc,
        ["Tabla", "Relaciones principales", "Uso"],
        [
            ["usuario", "usuario_grupo.id_usuario", "Registro de personas e integrantes."],
            ["grupo", "proceso.id_grupo, usuario_grupo.id_grupo", "Unidad responsable."],
            ["rol", "usuario_grupo.id_rol", "Papel dentro del grupo."],
            ["proceso", "riesgo_proceso.id_proceso", "Actividad donde se materializa el riesgo."],
            ["riesgo", "control.id_riesgo, riesgo_proceso.id_riesgo", "Evento evaluado por impacto/probabilidad."],
            ["control", "control.id_riesgo", "Medida que reduce exposicion."],
            ["riesgo_proceso", "id_riesgo + id_proceso", "Relacion muchos a muchos."],
        ],
        widths=[1.3, 2.5, 2.5],
    )

    add_heading(doc, "10. Recomendaciones tecnicas")
    add_number(doc, "Implementar login real con verificacion bcrypt.checkpw y control de sesion autenticada.", "C4")
    add_number(doc, "Mover secret_key, usuario, password, host, puerto y base de datos a variables de entorno.", "C2")
    add_number(doc, "Unificar creacion de tablas en migraciones para evitar DDL repartido dentro de rutas.", "C5")
    add_number(doc, "Agregar pruebas unitarias para calculo de nivel, riesgo inherente, residual y validacion de porcentajes.", "C6")
    add_number(doc, "Agregar control de integridad al eliminar riesgos con controles asociados, similar a la proteccion existente con procesos.", "C6")
    add_number(doc, "Considerar una libreria de Excel si se incrementa la complejidad de reportes, para reducir mantenimiento manual de XML.", "C10")

    add_heading(doc, "11. Conclusiones")
    add_paragraph(
        doc,
        "El sistema resuelve un flujo completo de gestion de riesgos: estructura organizacional, procesos, riesgos, controles, calculo inherente/residual, visualizacion ejecutiva y exportacion. La decision tecnica mas importante es que el dashboard no solo muestra datos, sino que integra metodologia de riesgo con responsables y controles, permitiendo priorizar acciones.",
    )
    add_paragraph(
        doc,
        "La base funcional es consistente, aunque quedan mejoras recomendadas en autenticacion, configuracion por entorno, migraciones y pruebas automatizadas. Esas mejoras no cambian el proposito del sistema: convertir registros operativos en informacion ejecutiva para decision.",
    )

    add_heading(doc, "12. Citas y fuentes del codigo")
    add_paragraph(doc, "Las citas siguientes corresponden a evidencias revisadas directamente en el repositorio local.")
    rows = [[key, value] for key, value in CITAS.items()]
    add_table(doc, ["Cita", "Fuente"], rows, widths=[0.75, 5.55])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("MAGERISK - Documentacion tecnica del sistema")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
